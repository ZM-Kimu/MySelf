import base64
import json
import os
import re
from dataclasses import asdict

import fitz
from docx import Document
from docx.document import Document as DocumentObject
from docx.oxml.ns import qn
from docx.oxml.text.run import CT_R
from docx.text.paragraph import Paragraph
from PIL import Image


def extract_image_from_docx(doc: DocumentObject, para: Paragraph) -> str:
    namespaces = {
        "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
        "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
        "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
    }

    for run in para.runs:
        elem: CT_R = getattr(run, "_element")
        drawing = elem.find(".//w:drawing", namespaces=elem.nsmap)
        if drawing is not None:
            blip = drawing.find(".//a:blip", namespaces=namespaces)
            if blip is not None:
                rId = blip.attrib.get(qn("r:embed"))
                image_part = doc.part.related_parts[rId]
                image_bytes = image_part.blob
                image_b64 = base64.b64encode(image_bytes).decode("utf-8")
                return image_b64

    return ""


def flatten_table(tables_data: list) -> str:
    string = ""
    for index, table in enumerate(tables_data):
        string += f"表格-{index+1}\n"
        for row in table:
            string += " ".join(row)
            string += "\n"
        string += "\n"

    return string


def extract_tables_from_docx(doc: DocumentObject) -> str:
    table_data = []

    for index, table in enumerate(doc.tables):
        table_data_per_table = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                cell_data = ""
                # 获取单元格内的所有文本内容
                for cel_para in cell.paragraphs:
                    for run in cel_para.runs:
                        cell_data += run.text.strip() + " "
                row_data.append(cell_data.strip())
            table_data_per_table.append(row_data)
        table_data.append(table_data_per_table)

    return flatten_table(table_data)


def extract_docx(file_path: str) -> str:
    doc = Document(file_path)
    doc_string = ""
    for para in doc.paragraphs:
        doc_string += para.text.strip()
        if img := extract_image_from_docx(doc, para):
            doc_string += f"<!image!>{img}<¡ǝƃɐɯı¡>"

        doc_string += "\n"

    if table_text := extract_tables_from_docx(doc):
        doc_string += table_text

    return doc_string


def extract_pdf(file_path: str) -> str:
    pdf = fitz.open(file_path)
    pdf_string = ""

    for page_num in range(len(pdf)):
        page = pdf.load_page(page_num)
        pdf_string += page.get_textpage().extractText()

    return pdf_string


if __name__ == "__main__":
    os.makedirs("questions", exist_ok=True)
    for file in os.listdir("./questions"):
        file_path = os.path.join("./questions", file)
        output_path = os.path.join("./questions", f"{os.path.splitext(file)[0]}.txt")

        full_text = ""
        if file_path.endswith("docx"):
            full_text = extract_docx(file_path)
        elif file_path.endswith("pdf"):
            full_text = extract_pdf(file_path)
        else:
            continue

        with open(output_path, "w", encoding="utf-8") as f:
            f.write(full_text)
